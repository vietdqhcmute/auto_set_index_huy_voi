import re
import os
import inspect
import win32com.client
from win32com.client import Dispatch

import docx.document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def format_path(_path):
    if "& " in _path or "'" in _path or '"' in _path:
        _path = _path.replace("'","")
        _path = _path.replace('"',"")
        _path = _path.replace('& ',"")
    return _path

#Add_paragraph_before_any_item
def insert_paragraph_before(item, text, style=None):
    """
    Return a newly created paragraph, inserted directly before this
    item (Table, etc.).
    """

    p = CT_P.add_p_before(item._element)
    p2 = Paragraph(p, item._parent)
    p2.text = text
    p2.style = style
    return p2

#Add_Table of Content
def add_table_of_content(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    # fldChar.set(qn('w:dirty'), 'true') #Make TOC auto update on openning
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p

def update_toc(file_name, pdf = "False"):
    script_dir = os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
    file_path = os.path.join(script_dir, file_name)
    word = win32com.client.DispatchEx("Word.Application")
    doc = word.Documents.Open(file_path)
    doc.TablesOfContents(1).Update()
    if pdf == "TRUE":
        doc.SaveAs2(file_name.replace(".docx", ".pdf"), FileFormat = 17)
    doc.Close(SaveChanges=True)
    word.Quit()

def create_run_xml(run):
    fldStart = OxmlElement('w:fldChar')
    fldStart.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'separate')

    fldChar2 = OxmlElement('w:t')
    fldChar2.text = "2"

    fldEnd = OxmlElement('w:fldChar')
    fldEnd.set(qn('w:fldCharType'), 'end')

    run._r.append(fldStart)
    run._r.append(instrText)
    run._r.append(fldChar1)
    run._r.append(fldChar2)
    run._r.append(fldEnd)

def add_page_number(doc):
    for section in doc.sections:
        create_run_xml(section.footer.add_paragraph().add_run())
        sectPr = section._sectPr

        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:start'), "0")
        sectPr.append(pgNumType)

def doc_to_docx(doc_path):
    word = Dispatch("Word.Application")
    word.visible = 0
    wb = word.Documents.Open(doc_path)
    out_file = doc_path.replace(".doc", ".docx")
    wb.SaveAs2(out_file, FileFormat=16) # file format for docx
    wb.Close()
    word.Quit()
    return out_file

def main():
    # doc_path = input("Nhập đường link file docx: ")
    doc_path ='/home/vietdo/code/auto_set_index_huy_voi/app/doc/3. Tai lieu Bo luat dan su.doc'
    doc_path = format_path(doc_path)
    _pdf = "TRUE"
    # _pdf = input("Nếu muốn tạo pdf thì gõ true: ")


    if doc_path.endswith(".doc"):
        doc_path = doc_to_docx(doc_path)

    doc = docx.Document(doc_path)
    styles = doc.styles


    ### Xử lý heading
    heading_list = ["Heading 1", "Heading 2", "Heading 3"]
    for heading in heading_list:
        if heading not in styles:
            styles.add_style(heading, WD_STYLE_TYPE.PARAGRAPH, builtin = True)

    key_words = ["(^Điều.*[.].*)", "^Mục.*[.]$"]
    key_words_2_line = ["^Chương.*", "^BỘ LUẬT.*", "^LUẬT.*", "^NGHỊ ĐỊNH.*", "^Phần thứ.*"]

    count = 1
    for i in range(0, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        for key_word in key_words:
            if re.match(key_word, para.text):
                para.style = styles["Heading 3"]
                para.style.font.size = Pt(14)
                print(count, para.text, para.style)
                count += 1
        for key_word in key_words_2_line:
            if re.match(key_word, para.text):
                para.style = styles["Heading 1"]
                para.style.font.size = Pt(16)
                doc.paragraphs[i+1].style = styles["Heading 2"]
                doc.paragraphs[i+1].style.font.size = Pt(16)
                print(count, para.text, para.style)
                count += 1

    #Thêm Mục lục
    for item in doc.iter_inner_content():
        para_clone = insert_paragraph_before(item, "")
        para_clone.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para_clone.style.font.size = Pt(12)
        add_table_of_content(para_clone)
        para_clone.add_run().add_break(WD_BREAK.PAGE)
        break

    #Thêm số trang
    add_page_number(doc)

    doc.save(doc_path)

    #Cập nhật mục lục và chuyển thành PDF
    update_toc(doc_path, _pdf.upper())

main()
